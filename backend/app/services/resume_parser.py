"""Resume parsing service for extracting text from PDF and DOCX files."""
import io
from typing import Optional
import PyPDF2
import pdfplumber
from docx import Document


class ResumeParser:
    """Service for parsing resume files and extracting text content."""
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """
        Extract text from a PDF file using multiple methods for better accuracy.
        
        Args:
            file_content: Binary content of the PDF file
            
        Returns:
            Extracted text content
        """
        text = ""
        
        # Try pdfplumber first (better for complex layouts)
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception:
            pass
        
        # Fallback to PyPDF2 if pdfplumber fails or returns empty
        if not text.strip():
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            except Exception as e:
                raise ValueError(f"Failed to parse PDF: {str(e)}")
        
        return text.strip()
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_content: Binary content of the DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            doc = Document(io.BytesIO(file_content))
            text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(" | ".join(row_text))
            
            return "\n".join(text)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    def parse_resume(file_content: bytes, filename: str) -> str:
        """
        Parse a resume file and extract text content.
        
        Args:
            file_content: Binary content of the file
            filename: Name of the file to determine format
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is not supported
        """
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return ResumeParser.extract_text_from_pdf(file_content)
        elif filename_lower.endswith('.docx'):
            return ResumeParser.extract_text_from_docx(file_content)
        elif filename_lower.endswith('.txt'):
            return file_content.decode('utf-8', errors='ignore')
        else:
            raise ValueError(
                f"Unsupported file format. Please upload PDF, DOCX, or TXT file. "
                f"Received: {filename}"
            )
    
    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Clean each line
            cleaned_line = ' '.join(line.split())
            if cleaned_line:
                cleaned_lines.append(cleaned_line)
        
        return '\n'.join(cleaned_lines)


# Create singleton instance
resume_parser = ResumeParser()
